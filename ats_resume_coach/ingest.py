"""Input loading and document text extraction."""

from __future__ import annotations

import ipaddress
import socket
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse


TEXT_SUFFIXES = {".txt", ".md", ".rst"}
HTML_SUFFIXES = {".html", ".htm"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}

FETCH_HEADER_PROFILES = (
    {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Cache-Control": "no-cache",
        "Pragma": "no-cache",
        "Upgrade-Insecure-Requests": "1",
        "Sec-Fetch-Dest": "document",
        "Sec-Fetch-Mode": "navigate",
        "Sec-Fetch-Site": "none",
        "Sec-Fetch-User": "?1",
    },
    {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.5 Safari/605.1.15",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.8",
    },
)


class IngestError(ValueError):
    """Raised when a document cannot be read or safely fetched."""


def read_file_text(path: str | Path) -> str:
    file_path = Path(path)
    if not file_path.exists():
        raise IngestError(f"File does not exist: {file_path}")
    return extract_text_from_bytes(file_path.read_bytes(), filename=file_path.name)


def extract_text_from_bytes(content: bytes, *, filename: str = "", content_type: str = "") -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in PDF_SUFFIXES or "pdf" in content_type:
        return _extract_pdf(content)
    if suffix in DOCX_SUFFIXES or "wordprocessingml" in content_type:
        return _extract_docx(content)
    if suffix in HTML_SUFFIXES or "html" in content_type:
        return _extract_html(content)
    return _decode_text(content)


def fetch_url_text(url: str, *, timeout_seconds: int = 12) -> str:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise IngestError("Only public http(s) URLs are supported.")
    _ensure_public_hostname(parsed.hostname or "")

    import requests
    last_error: Exception | None = None

    for headers in FETCH_HEADER_PROFILES:
        try:
            response = requests.get(url, timeout=timeout_seconds, headers=headers)
            if response.status_code >= 400:
                last_error = requests.HTTPError(
                    f"{response.status_code} Client Error: {response.reason} for url: {url}",
                    response=response,
                )
                continue
            content_type = response.headers.get("content-type", "")
            return extract_text_from_bytes(response.content, filename=parsed.path, content_type=content_type)
        except requests.RequestException as exc:
            last_error = exc

    message = "Could not fetch the job URL. Some sites block automated requests; paste the job text or upload the posting instead."
    if last_error is not None:
        raise IngestError(message) from last_error
    raise IngestError(message)


def _ensure_public_hostname(hostname: str) -> None:
    if not hostname:
        raise IngestError("URL hostname is missing.")
    try:
        addresses = socket.getaddrinfo(hostname, None)
    except socket.gaierror as exc:
        raise IngestError(f"Could not resolve URL hostname: {hostname}") from exc

    for address in addresses:
        ip = ipaddress.ip_address(address[4][0])
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved or ip.is_multicast:
            raise IngestError("Private, local, or reserved network URLs are not allowed.")


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
        from pypdf.errors import PyPdfError
    except ImportError as exc:
        raise IngestError("PDF parsing requires the pypdf dependency.") from exc

    try:
        reader = PdfReader(BytesIO(content))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
    except PyPdfError as exc:
        raise IngestError("Could not extract text from PDF.") from exc
    return "\n".join(pages).strip()


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise IngestError("DOCX parsing requires the python-docx dependency.") from exc

    try:
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    value = cell.text.strip()
                    if value:
                        table_cells.append(value)
    except Exception as exc:
        raise IngestError("Could not extract text from DOCX.") from exc
    return "\n".join(paragraphs + table_cells).strip()


def _extract_html(content: bytes) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise IngestError("HTML parsing requires the beautifulsoup4 dependency.") from exc

    soup = BeautifulSoup(_decode_text(content), "html.parser")
    for tag in soup(["script", "style", "noscript", "svg"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise IngestError("Could not decode document text.")
