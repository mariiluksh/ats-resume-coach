"""Build a tailored resume draft from the analysis output."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.document import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from .schemas import AnalysisResult
from .taxonomy import SKILL_ALIASES
from .text import (
    contains_phrase,
    extract_lines,
    has_email,
    has_phone,
    normalize_text,
    normalize_for_match,
)


@dataclass(frozen=True)
class DraftResult:
    filename: str
    content_type: str
    data: bytes
    preview_text: str


@dataclass(frozen=True)
class SourceLine:
    text: str
    style_name: str
    is_bullet: bool = False


@dataclass(frozen=True)
class SourceSection:
    heading: str
    style_name: str
    lines: list[SourceLine]


@dataclass(frozen=True)
class ParsedResume:
    header: list[SourceLine]
    sections: list[SourceSection]


@dataclass(frozen=True)
class ResumeStyles:
    name: str
    contact: str
    section_heading: str
    item_heading: str
    body: str
    bullet: str


DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SECTION_HEADINGS = {
    "education",
    "experience",
    "work experience",
    "professional experience",
    "employment",
    "projects",
    "technical projects",
    "skills",
    "technical skills",
    "core skills",
    "summary",
    "profile",
    "professional summary",
    "selected projects",
    "selected projects and research",
    "selected projects & research",
    "certifications",
    "awards",
    "leadership and awards",
    "leadership & awards",
    "leadership",
    "publications",
    "interests",
    "professional interests",
    "product interests",
    "languages",
}
SUMMARY_HEADINGS = {"summary", "profile", "professional summary"}
SKILLS_HEADINGS = {"skills", "technical skills", "core skills"}
PROJECT_HEADINGS = {"projects", "technical projects", "selected projects", "selected projects and research", "selected projects & research"}
INTEREST_HEADINGS = {"interests", "professional interests", "product interests"}
LEADERSHIP_HEADINGS = {"leadership", "awards", "leadership and awards", "leadership & awards"}
LANGUAGE_HEADINGS = {"languages"}
EXPERIENCE_HEADINGS = {"experience", "work experience", "professional experience", "employment"}
EDUCATION_HEADINGS = {"education"}
TARGET_SECTION_ORDER = (
    "PROFILE",
    "EDUCATION",
    "EXPERIENCE",
    "SELECTED PROJECTS & RESEARCH",
    "LEADERSHIP & AWARDS",
    "TECHNICAL SKILLS",
    "PROFESSIONAL INTERESTS",
    "LANGUAGES",
)
SECTION_KIND_TO_HEADING = {
    "profile": "PROFILE",
    "education": "EDUCATION",
    "experience": "EXPERIENCE",
    "projects": "SELECTED PROJECTS & RESEARCH",
    "leadership": "LEADERSHIP & AWARDS",
    "skills": "TECHNICAL SKILLS",
    "interests": "PROFESSIONAL INTERESTS",
    "languages": "LANGUAGES",
}
FINANCE_TERMS = (
    "investment management",
    "financial markets",
    "quantitative finance",
    "quantitative trading",
    "trading",
    "asset management",
    "hedge funds",
    "fintech",
    "financial technology",
)
TECH_TERMS = (
    "artificial intelligence",
    "machine learning",
    "software engineering",
    "data analysis",
    "frontend",
    "full-stack",
    "api",
)
DATE_RE = re.compile(
    r"\b(?:20\d{2}|19\d{2}|present|current|jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec|summer|spring|autumn|fall)\b",
    re.IGNORECASE,
)
ROLE_SIGNAL_RE = re.compile(
    r"\b(?:engineer|developer|analyst|assistant|intern|participant|programme|program|secretary|representative|researcher|consultant|manager|lead|officer)\b",
    re.IGNORECASE,
)
ACTION_START_RE = re.compile(
    r"^(?:built|build|created|create|developed|develop|implemented|implement|improved|improve|reduced|reduce|increase|increased|"
    r"analyzed|analysed|analyze|analyse|automated|automate|designed|design|deployed|deploy|tested|test|documented|document|"
    r"collaborated|collaborate|presented|present|optimized|optimised|optimize|optimise|led|lead|managed|manage|selected|select|"
    r"completed|complete|worked|work|performed|perform|conducted|conduct|contributed|contribute|secured|secure|supported|support|"
    r"queried|query|validated|validate|evaluated|evaluate|researched|research|explored|explore|ensured|ensure|engaged|engage|"
    r"attended|attend|identified|identify|used|use)\b",
    re.IGNORECASE,
)
SKILL_LABELS = {
    "programming": "Programming",
    "data": "Data & Analytics",
    "web": "Technologies",
    "tools": "Tools",
    "cloud": "Cloud",
    "professional": "Professional",
}


def build_resume_draft(
    job_text: str,
    resume_text: str,
    analysis: AnalysisResult,
    *,
    source_docx: bytes | None = None,
    source_filename: str | None = None,
) -> DraftResult:
    job_text = normalize_text(job_text)
    resume_text = _merge_wrapped_lines(normalize_text(resume_text))
    if source_docx:
        return _build_style_preserving_docx(
            source_docx=source_docx,
            source_filename=source_filename,
            job_text=job_text,
            resume_text=resume_text,
            analysis=analysis,
        )

    lines = extract_lines(resume_text)

    draft_lines = _build_text_draft(lines, job_text, resume_text, analysis)
    draft_text = "\n".join(draft_lines).strip() + "\n"

    document = Document()
    _apply_resume_document_formatting(document)
    for section in _build_doc_sections(draft_lines):
        heading, entries = section
        if heading == "Header":
            for index, entry in enumerate(entries):
                role = "name" if index == 0 else "contact"
                _add_doc_paragraph(document, entry, "Normal", role=role)
            continue
        if heading != "Header":
            _add_doc_paragraph(document, heading, "Heading 1", role="section_heading")
        for entry in entries:
            if entry.startswith("- "):
                _add_doc_paragraph(document, entry[2:], "List Bullet", role="bullet")
            else:
                _add_doc_paragraph(document, entry, "Normal", role="body")

    buffer = BytesIO()
    document.save(buffer)

    filename = _draft_filename(analysis)
    return DraftResult(
        filename=_source_or_draft_filename(source_filename, analysis),
        content_type=DOCX_CONTENT_TYPE,
        data=buffer.getvalue(),
        preview_text=draft_text,
    )


def build_resume_draft_preview(job_text: str, resume_text: str, analysis: AnalysisResult) -> str:
    return build_resume_draft(job_text, resume_text, analysis).preview_text


def _build_text_draft(
    lines: list[str],
    job_text: str,
    resume_text: str,
    analysis: AnalysisResult,
) -> list[str]:
    merged_lines = _merge_wrapped_lines_from_lines(lines)
    header = _build_header(merged_lines, resume_text)
    parsed = _parse_text_resume(merged_lines, header)
    summary = _build_profile(parsed, analysis, job_text, resume_text)
    skill_lines = _build_skill_lines(analysis, resume_text)
    sections: list[str] = []
    sections.extend(header)
    sections.extend(["", "PROFILE", summary])

    for output_heading in TARGET_SECTION_ORDER:
        if output_heading == "PROFILE":
            continue
        if output_heading == "TECHNICAL SKILLS":
            if skill_lines:
                sections.extend(["", output_heading])
                sections.extend(skill_lines)
            continue
        source_sections = _sections_for_output_heading(parsed, output_heading)
        if not source_sections and output_heading == "PROFESSIONAL INTERESTS":
            interest_line = _build_interest_line(job_text, resume_text)
            if interest_line:
                sections.extend(["", output_heading, interest_line])
            continue
        if not source_sections:
            continue
        sections.extend(["", output_heading])
        for source_section in source_sections:
            sections.extend(_section_text_lines(source_section, output_heading, analysis, job_text, resume_text))

    return sections


def _build_header(lines: list[str], resume_text: str) -> list[str]:
    candidates = []
    for line in lines[:5]:
        if has_email(line) or has_phone(line):
            candidates.append(line)
    if not candidates and lines:
        candidates.append(lines[0])
        candidates.extend([line for line in lines[1:5] if has_email(line) or has_phone(line)])
    if not candidates:
        candidates.append("Name")
    return candidates


def _build_summary(analysis: AnalysisResult) -> str:
    return _build_profile(ParsedResume(header=[], sections=[]), analysis, "", "")


def _build_skills(analysis: AnalysisResult) -> str:
    return " | ".join(_build_skill_lines(analysis, ""))


def _build_style_preserving_docx(
    *,
    source_docx: bytes,
    source_filename: str | None,
    job_text: str,
    resume_text: str,
    analysis: AnalysisResult,
) -> DraftResult:
    document = Document(BytesIO(source_docx))
    parsed = _parse_docx_resume(document)
    summary = _build_profile(parsed, analysis, job_text, resume_text)
    preview_lines = ["Preserved source DOCX structure with tailored content.", "PROFILE", summary]
    _rewrite_docx_in_place(document, parsed, analysis, job_text, resume_text, summary)

    buffer = BytesIO()
    document.save(buffer)

    return DraftResult(
        filename=_source_or_draft_filename(source_filename, analysis),
        content_type=DOCX_CONTENT_TYPE,
        data=buffer.getvalue(),
        preview_text="\n".join(preview_lines).strip() + "\n",
    )


def _rewrite_docx_in_place(
    document: DocxDocument,
    parsed: ParsedResume,
    analysis: AnalysisResult,
    job_text: str,
    resume_text: str,
    summary: str,
) -> None:
    del parsed  # structure is preserved by paragraph order; parsing is only used for summary generation
    current_section = ""
    skills_lines = _build_skill_lines(analysis, resume_text)
    skills_text = " | ".join(skills_lines)
    interest_line = _build_interest_line(job_text, resume_text)
    summary_written = False
    skills_written = False
    interest_written = False

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        normalized = _normalize_heading(text)
        if normalized in SUMMARY_HEADINGS:
            current_section = "PROFILE"
            continue
        if normalized in SKILLS_HEADINGS:
            current_section = "SKILLS"
            continue
        if normalized in INTEREST_HEADINGS:
            current_section = "INTERESTS"
            continue
        if normalized in EXPERIENCE_HEADINGS:
            current_section = "EXPERIENCE"
            continue
        if normalized in PROJECT_HEADINGS:
            current_section = "PROJECTS"
            continue
        if normalized in EDUCATION_HEADINGS:
            current_section = "EDUCATION"
            continue
        if normalized in LEADERSHIP_HEADINGS:
            current_section = "LEADERSHIP"
            continue
        if normalized in LANGUAGE_HEADINGS:
            current_section = "LANGUAGES"
            continue

        if current_section == "PROFILE" and not summary_written and text:
            _set_paragraph_text(paragraph, summary)
            summary_written = True
            continue
        if current_section == "SKILLS" and not skills_written and text:
            _set_paragraph_text(paragraph, skills_text or text)
            skills_written = True
            continue
        if current_section == "INTERESTS" and not interest_written and text:
            _set_paragraph_text(paragraph, interest_line or text)
            interest_written = True
            continue
        if current_section in {"EXPERIENCE", "PROJECTS"} and text:
            if _looks_like_bullet_content(text) or paragraph.style.name.startswith("List"):
                _set_paragraph_text(paragraph, _polish_bullet(text, analysis))
            elif _looks_like_role_fragment(text) or _is_role_or_date_line(SourceLine(text=text, style_name=paragraph.style.name)):
                _set_paragraph_text(paragraph, _clean_heading_line(text))
            continue
        if current_section == "EDUCATION" and text:
            continue
        if current_section == "LEADERSHIP" and text:
            continue


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.add_run(text)


def _parse_docx_resume(document: DocxDocument) -> ParsedResume:
    header: list[SourceLine] = []
    sections: list[SourceSection] = []
    current: SourceSection | None = None

    for paragraph in document.paragraphs:
        for text, is_heading in _split_docx_paragraph_text(paragraph.text):
            text = normalize_text(text)
            if not text:
                continue
            line = SourceLine(text=text, style_name=paragraph.style.name, is_bullet=_is_bullet_paragraph(paragraph))
            if is_heading or _looks_like_heading(text):
                current = SourceSection(heading=_display_heading(text), style_name=paragraph.style.name, lines=[])
                sections.append(current)
                continue
            if current is None:
                header.append(line)
            else:
                current.lines.append(line)

    return ParsedResume(header=header, sections=sections)


def _split_docx_paragraph_text(text: str) -> list[tuple[str, bool]]:
    cleaned = normalize_text(text)
    if not cleaned:
        return []
    heading_names = sorted(
        {
            *SECTION_HEADINGS,
            *SUMMARY_HEADINGS,
            *SKILLS_HEADINGS,
            *PROJECT_HEADINGS,
            *INTEREST_HEADINGS,
            *LEADERSHIP_HEADINGS,
            *LANGUAGE_HEADINGS,
        },
        key=len,
        reverse=True,
    )
    pattern = re.compile(
        "|".join(rf"(?<!\w){re.escape(name)}(?!\w)" for name in heading_names),
        re.IGNORECASE,
    )
    parts: list[tuple[str, bool]] = []
    last = 0
    for match in pattern.finditer(cleaned):
        before = cleaned[last:match.start()].strip(" -:\n\t")
        if before:
            parts.append((before, False))
        parts.append((match.group(0).strip(), True))
        last = match.end()
    tail = cleaned[last:].strip(" -:\n\t")
    if tail:
        parts.append((tail, False))
    return parts or [(cleaned, False)]


def _parse_text_resume(lines: list[str], header: list[str]) -> ParsedResume:
    parsed_header = [SourceLine(text=line, style_name="Normal") for line in header]
    sections: list[SourceSection] = []
    current: SourceSection | None = None
    header_values = set(header)

    for line in lines:
        if line in header_values:
            continue
        if _looks_like_heading(line):
            current = SourceSection(heading=_display_heading(line), style_name="Heading 1", lines=[])
            sections.append(current)
            continue
        if current is None:
            current = SourceSection(heading="EXPERIENCE", style_name="Heading 1", lines=[])
            sections.append(current)
        current.lines.append(SourceLine(text=_format_source_line(line), style_name="Normal", is_bullet=line.startswith(("-", "*", "•"))))

    return ParsedResume(header=parsed_header, sections=sections)


def _merge_wrapped_lines_from_lines(lines: list[str]) -> list[str]:
    return _merge_wrapped_lines("\n".join(lines)).splitlines()


def _merge_wrapped_lines(text: str) -> str:
    lines = extract_lines(text)
    merged: list[str] = []
    for line in lines:
        if not merged:
            merged.append(line)
            continue
        previous = merged[-1]
        if _should_merge_wrapped_line(previous, line):
            merged[-1] = f"{previous} {line}".strip()
        else:
            merged.append(line)
    return "\n".join(merged)


def _should_merge_wrapped_line(previous: str, current: str) -> bool:
    if _looks_like_heading(previous) or _looks_like_heading(current):
        return False
    if "@" in previous or "@" in current:
        return False
    if has_email(previous) or has_email(current) or has_phone(previous) or has_phone(current):
        return False
    if "linkedin.com" in normalize_for_match(previous) or "linkedin.com" in normalize_for_match(current):
        return False
    if "github.com" in normalize_for_match(previous) or "github.com" in normalize_for_match(current):
        return False
    if previous.startswith(("•", "-", "*")) or current.startswith(("•", "-", "*")):
        return False
    if previous.endswith((".", "!", "?", ":", ";")):
        return False
    if current[:1].islower():
        return True
    if len(previous.split()) <= 4 or len(current.split()) <= 4:
        return False
    if len(previous.split()) <= 12 and len(current.split()) <= 10:
        return True
    return False


def _resume_styles(document: DocxDocument, parsed: ParsedResume) -> ResumeStyles:
    header_styles = [line.style_name for line in parsed.header]
    section_style = parsed.sections[0].style_name if parsed.sections else "Heading 1"
    item_heading_style = _first_line_style(parsed, style_signal="Heading 3") or _first_line_style(parsed) or "Normal"
    body_style = _first_body_style(parsed) or "Normal"
    bullet_style = "List Bullet" if _style_exists(document, "List Bullet") else body_style
    return ResumeStyles(
        name=header_styles[0] if header_styles else "Title",
        contact=header_styles[1] if len(header_styles) > 1 else body_style,
        section_heading=section_style,
        item_heading=item_heading_style,
        body=body_style,
        bullet=bullet_style,
    )


def _first_line_style(parsed: ParsedResume, style_signal: str | None = None) -> str | None:
    for section in parsed.sections:
        for line in section.lines:
            if style_signal is None or style_signal.lower() in line.style_name.lower():
                return line.style_name
    return None


def _first_body_style(parsed: ParsedResume) -> str | None:
    for section in parsed.sections:
        for line in section.lines:
            if "heading" not in line.style_name.lower():
                return line.style_name
    return None


def _build_profile(parsed: ParsedResume, analysis: AnalysisResult, job_text: str, resume_text: str) -> str:
    education = _education_profile_fragment(parsed)
    experience = _experience_profile_fragment(parsed, resume_text)
    leadership = _leadership_profile_fragment(resume_text)
    strengths = _strength_terms(analysis, resume_text)
    interests = _domain_terms(resume_text)
    role = _clean_role(analysis.job_title_guess)

    sentences: list[str] = []
    opener_parts: list[str] = []
    if education:
        opener_parts.append(education)
    else:
        opener_parts.append(f"Candidate for {role}" if role else "Candidate")
    if experience:
        opener_parts.append(experience)
    sentences.append(" ".join(opener_parts).strip().rstrip(".") + ".")

    if strengths:
        sentences.append("Strong in " + _join_human(strengths[:5]) + ".")
    if leadership:
        sentences.append(leadership)
    if interests:
        sentences.append("Interested in " + _join_human(interests[:5]) + ".")
    elif role:
        sentences.append(f"Focused on applying technical and analytical skills to {role.lower()} roles.")

    return " ".join(_dedupe_sentences(sentences))


def _education_profile_fragment(parsed: ParsedResume) -> str | None:
    education = _first_section(parsed, "EDUCATION")
    if not education:
        return None
    degree = None
    institution = None
    for line in education.lines:
        degree_match = re.search(
            r"\b(?:MEng|BSc|BA|MSc|BS|MS|PhD|Bachelor(?:'s)?|Master(?:'s)?)[^,.;|]{0,90}",
            line.text,
            re.IGNORECASE,
        )
        if degree_match and degree is None:
            degree = degree_match.group(0).strip()
        lowered = line.text.lower()
        if institution is None and any(marker in lowered for marker in ("university", "college", "school")):
            if "—" in line.text or " - " in line.text:
                institution = re.split(r"\s+[—|-]\s+", line.text, maxsplit=1)[-1].strip()
            elif "," in line.text:
                institution = line.text.rsplit(",", 1)[-1].strip()
            else:
                institution = line.text.strip()
    if degree and institution:
        return f"{degree} student at {institution}"
    if degree:
        return f"{degree} student"
    return None


def _experience_profile_fragment(parsed: ParsedResume, resume_text: str) -> str | None:
    experience = _first_section(parsed, "EXPERIENCE")
    if not experience:
        return None

    lowered = normalize_for_match(resume_text)
    if any(signal in lowered for signal in ("software engineer", "developer", "frontend")) and any(
        signal in lowered for signal in ("trading", "platform", "financial", "market", "portfolio")
    ):
        if "frontend" in lowered:
            return "with commercial software engineering experience developing frontend trading technology for financial markets"
        return "with commercial software engineering experience developing trading technology for financial markets"
    if any(signal in lowered for signal in ("data analysis", "data analytics", "analytics", "sql", "pandas", "numpy")):
        return "with experience building data-driven tools and analytical workflows"
    if any(signal in lowered for signal in ("machine learning", "research", "model", "algorithm")):
        return "with experience applying analytical and research methods to technical projects"

    role_line = None
    context_line = None
    for line in experience.lines[:8]:
        if role_line is None and ROLE_SIGNAL_RE.search(line.text):
            role_line = _remove_date_suffix(line.text)
        elif context_line is None and any(signal in line.text.lower() for signal in ("platform", "software", "data", "financial", "trading", "product")):
            context_line = line.text
    if role_line and context_line:
        context = _compact_context(context_line)
        return f"with experience as {_role_with_article(role_line)} {context}"
    if role_line:
        return f"with experience as {_role_with_article(role_line)}"
    return None


def _leadership_profile_fragment(resume_text: str) -> str | None:
    lowered = resume_text.lower()
    if any(signal in lowered for signal in ("executive committee", "society", "student representative", "partnership", "outreach")):
        return "Brings leadership, stakeholder engagement and communication experience across student and professional environments."
    return None


def _build_skill_lines(analysis: AnalysisResult, resume_text: str) -> list[str]:
    grouped: dict[str, list[str]] = {label: [] for label in SKILL_LABELS.values()}
    for category, values in _resume_skill_matches(resume_text, analysis).items():
        label = SKILL_LABELS.get(category, category.title())
        grouped.setdefault(label, [])
        grouped[label].extend(_title_skill(value) for value in values)

    for label, values in _existing_skill_groups(resume_text).items():
        grouped.setdefault(label, [])
        grouped[label].extend(values)

    lines = []
    for label in ("Programming", "Technologies", "Data & Analytics", "Tools", "Cloud", "Professional"):
        values = _dedupe([value for value in grouped.get(label, []) if value])
        if values:
            lines.append(f"{label}: {' • '.join(values[:10])}")
    if not lines and analysis.matched_keywords:
        lines.append("Relevant Keywords: " + " • ".join(_title_skill(value) for value in analysis.matched_keywords[:8]))
    return lines


def _resume_skill_matches(resume_text: str, analysis: AnalysisResult) -> dict[str, list[str]]:
    matches: dict[str, list[str]] = {}
    for category, skills in SKILL_ALIASES.items():
        values = set(analysis.matched_skills.get(category, []))
        for skill, aliases in skills.items():
            if any(_skill_alias_present(resume_text, alias) for alias in aliases):
                values.add(skill)
        if values:
            matches[category] = sorted(values)
    return matches


def _skill_alias_present(text: str, alias: str) -> bool:
    if alias.strip().lower() in {"java", "go", "r", "c"}:
        return bool(re.search(rf"\b{re.escape(alias.strip())}\b", text, flags=re.IGNORECASE))
    return contains_phrase(text, alias)


def _existing_skill_groups(resume_text: str) -> dict[str, list[str]]:
    lines = extract_lines(resume_text)
    groups: dict[str, list[str]] = {}
    active = False
    for line in lines:
        normalized = _normalize_heading(line)
        if normalized in SKILLS_HEADINGS:
            active = True
            continue
        if active and _looks_like_heading(line):
            break
        if not active:
            continue
        for label, raw_values in re.findall(r"([A-Za-z][A-Za-z &/]+):\s*([^:]+?)(?=\s+[A-Za-z][A-Za-z &/]+:|$)", line):
            cleaned_label = _skill_group_label(label)
            values = [_title_skill(value) for value in re.split(r"[,•|;]", raw_values) if value.strip()]
            groups.setdefault(cleaned_label, [])
            groups[cleaned_label].extend(values)
    return groups


def _skill_group_label(label: str) -> str:
    normalized = label.strip().lower()
    if "program" in normalized:
        return "Programming"
    if any(word in normalized for word in ("data", "analytics")):
        return "Data & Analytics"
    if any(word in normalized for word in ("frontend", "technology", "technologies", "web")):
        return "Technologies"
    if "cloud" in normalized:
        return "Cloud"
    if any(word in normalized for word in ("tool", "other")):
        return "Tools"
    return label.strip().title()


def _sections_for_output_heading(parsed: ParsedResume, output_heading: str) -> list[SourceSection]:
    return [section for section in parsed.sections if _canonical_output_heading(section.heading) == output_heading]


def _first_section(parsed: ParsedResume, output_heading: str) -> SourceSection | None:
    matches = _sections_for_output_heading(parsed, output_heading)
    return matches[0] if matches else None


def _canonical_output_heading(heading: str) -> str:
    normalized = _normalize_heading(heading)
    if normalized in SUMMARY_HEADINGS:
        return SECTION_KIND_TO_HEADING["profile"]
    if normalized in EDUCATION_HEADINGS:
        return SECTION_KIND_TO_HEADING["education"]
    if normalized in EXPERIENCE_HEADINGS:
        return SECTION_KIND_TO_HEADING["experience"]
    if normalized in PROJECT_HEADINGS:
        return SECTION_KIND_TO_HEADING["projects"]
    if normalized in LEADERSHIP_HEADINGS:
        return SECTION_KIND_TO_HEADING["leadership"]
    if normalized in SKILLS_HEADINGS:
        return SECTION_KIND_TO_HEADING["skills"]
    if normalized in INTEREST_HEADINGS:
        return SECTION_KIND_TO_HEADING["interests"]
    if normalized in LANGUAGE_HEADINGS:
        return SECTION_KIND_TO_HEADING["languages"]
    return heading.upper()


def _section_text_lines(
    section: SourceSection,
    output_heading: str,
    analysis: AnalysisResult,
    job_text: str,
    resume_text: str,
) -> list[str]:
    styles = ResumeStyles("Normal", "Normal", "Heading 1", "Normal", "Normal", "List Bullet")
    output_lines = _section_docx_lines(section, output_heading, styles, analysis, job_text, resume_text)
    return [f"- {line.text}" if line.is_bullet else line.text for line in output_lines]


def _section_docx_lines(
    section: SourceSection,
    output_heading: str,
    styles: ResumeStyles,
    analysis: AnalysisResult,
    job_text: str,
    resume_text: str,
) -> list[SourceLine]:
    if output_heading == "EXPERIENCE":
        return _polished_experience_lines(section, styles, analysis)
    if output_heading == "SELECTED PROJECTS & RESEARCH":
        return _polished_project_lines(section, styles, analysis, resume_text)
    if output_heading == "PROFESSIONAL INTERESTS":
        interest_line = _build_interest_line(job_text, resume_text)
        if interest_line:
            return [SourceLine(interest_line, styles.body)]
    if output_heading == "TECHNICAL SKILLS":
        return [SourceLine(line, styles.body) for line in _build_skill_lines(analysis, resume_text)]
    return [_normalize_source_line(line, styles) for line in section.lines]


def _polished_experience_lines(section: SourceSection, styles: ResumeStyles, analysis: AnalysisResult) -> list[SourceLine]:
    result: list[SourceLine] = []
    for index, line in enumerate(section.lines):
        previous = section.lines[index - 1] if index > 0 else None
        next_line = section.lines[index + 1] if index + 1 < len(section.lines) else None
        if _is_item_heading_line(line, previous, next_line):
            result.append(SourceLine(_clean_heading_line(line.text), line.style_name or styles.item_heading))
        elif _is_role_or_date_line(line) or _looks_like_role_fragment(line.text) or not _looks_like_bullet_content(line.text):
            result.append(SourceLine(_clean_heading_line(line.text), line.style_name or styles.body))
        else:
            result.append(SourceLine(_polish_bullet(line.text, analysis), styles.bullet, is_bullet=True))
    return result


def _polished_project_lines(
    section: SourceSection,
    styles: ResumeStyles,
    analysis: AnalysisResult,
    resume_text: str,
) -> list[SourceLine]:
    result: list[SourceLine] = []
    for index, line in enumerate(section.lines):
        previous = section.lines[index - 1] if index > 0 else None
        next_line = section.lines[index + 1] if index + 1 < len(section.lines) else None
        if _is_item_heading_line(line, previous, next_line):
            result.append(SourceLine(_clean_heading_line(line.text), line.style_name or styles.item_heading))
        elif _looks_like_role_fragment(line.text) or not _looks_like_bullet_content(line.text):
            result.append(SourceLine(_clean_heading_line(line.text), line.style_name or styles.body))
        else:
            result.append(SourceLine(_polish_bullet(line.text, analysis), styles.bullet, is_bullet=True))

    for research_line in _research_lines_from_resume(resume_text):
        if not any(research_line.lower() in line.text.lower() for line in result):
            result.append(SourceLine(research_line, styles.body))
    return result


def _normalize_source_line(line: SourceLine, styles: ResumeStyles) -> SourceLine:
    if line.is_bullet or _looks_like_bullet_content(line.text):
        return SourceLine(_polish_bullet(line.text, None), styles.bullet, is_bullet=True)
    return SourceLine(_clean_heading_line(line.text), line.style_name or styles.body)


def _build_doc_sections(lines: list[str]) -> list[tuple[str, list[str]]]:
    sections: list[tuple[str, list[str]]] = []
    current_heading = "Header"
    current_entries: list[str] = []
    for line in lines:
        if _is_output_section_heading(line):
            if current_entries:
                sections.append((current_heading, current_entries))
            current_heading = line
            current_entries = []
            continue
        current_entries.append(line)
    if current_entries:
        sections.append((current_heading, current_entries))
    return sections


def _draft_filename(analysis: AnalysisResult) -> str:
    slug = "tailored_resume"
    if analysis.job_title_guess:
        slug = "".join(ch.lower() if ch.isalnum() else "_" for ch in analysis.job_title_guess)[:40].strip("_") or slug
    return f"{slug}.docx"


def _tailored_filename(source_filename: str | None, analysis: AnalysisResult) -> str:
    if source_filename:
        stem = Path(source_filename).stem
        cleaned = "".join(ch if ch.isalnum() or ch in {" ", "-", "_"} else "_" for ch in stem).strip()
        if cleaned:
            return f"{cleaned}_tailored.docx"
    return _draft_filename(analysis)


def _source_or_draft_filename(source_filename: str | None, analysis: AnalysisResult) -> str:
    if source_filename:
        stem = Path(source_filename).stem
        cleaned = "".join(ch if ch.isalnum() or ch in {" ", "-", "_"} else "_" for ch in stem).strip()
        if cleaned:
            return f"{cleaned}_tailored.docx"
    return _draft_filename(analysis)


def _normalize_heading(text: str) -> str:
    return " ".join(text.lower().strip().strip(":").split())


def _looks_like_heading(line: str) -> bool:
    normalized = _normalize_heading(line)
    if normalized in SECTION_HEADINGS:
        return True
    words = normalized.split()
    return 1 <= len(words) <= 4 and line.upper() == line and any(word in SECTION_HEADINGS for word in words)


def _is_output_section_heading(line: str) -> bool:
    return line in {"Professional Summary", "Core Skills"} or _normalize_heading(line) in SECTION_HEADINGS


def _display_heading(line: str) -> str:
    normalized = _normalize_heading(line)
    if normalized in {"technical skills", "skills"}:
        return "Core Skills"
    if normalized in SUMMARY_HEADINGS:
        return "Professional Summary"
    return line.strip().strip(":").title()


def _format_source_line(line: str) -> str:
    cleaned = line.strip()
    if cleaned.startswith(("•", "-", "*")):
        return "- " + cleaned.lstrip("•-* \t")
    return cleaned


def _clear_document_body(document: DocxDocument) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _write_header(document: DocxDocument, parsed: ParsedResume, styles: ResumeStyles) -> None:
    for index, line in enumerate(parsed.header):
        style = styles.name if index == 0 else styles.contact
        role = "name" if index == 0 else "contact"
        _add_doc_paragraph(document, line.text, style, role=role)


def _write_section(document: DocxDocument, heading: str, lines: list[SourceLine], styles: ResumeStyles) -> None:
    if not lines:
        return
    _add_doc_paragraph(document, heading, styles.section_heading, role="section_heading")
    for line in lines:
        style = styles.bullet if line.is_bullet else line.style_name or styles.body
        role = "bullet" if line.is_bullet else ("item_heading" if style == styles.item_heading else "body")
        _add_doc_paragraph(document, line.text, style, role=role)


def _add_doc_paragraph(document: DocxDocument, text: str, style_name: str | None, *, role: str = "body") -> Paragraph:
    paragraph = document.add_paragraph()
    _apply_style(paragraph, style_name)
    run = paragraph.add_run(text)
    _format_doc_paragraph(paragraph, run, role=role)
    return paragraph


def _apply_style(paragraph: Paragraph, style_name: str | None) -> None:
    if not style_name:
        return
    try:
        paragraph.style = style_name
    except (KeyError, ValueError):
        return


def _style_exists(document: DocxDocument, style_name: str) -> bool:
    try:
        document.styles[style_name]
    except KeyError:
        return False
    return True


def _apply_resume_document_formatting(document: DocxDocument) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    _set_style_font(document, "Normal", name="Aptos", size=Pt(10))
    _set_style_font(document, "Title", name="Aptos", size=Pt(14), bold=True)
    _set_style_font(document, "Heading 1", name="Aptos", size=Pt(10.5), bold=True)
    if _style_exists(document, "List Bullet"):
        _set_style_font(document, "List Bullet", name="Aptos", size=Pt(10))


def _set_style_font(
    document: DocxDocument,
    style_name: str,
    *,
    name: str | None = None,
    size: Pt | None = None,
    bold: bool | None = None,
) -> None:
    try:
        style = document.styles[style_name]
    except KeyError:
        return
    if name:
        style.font.name = name
    if size is not None:
        style.font.size = size
    if bold is not None:
        style.font.bold = bold


def _format_doc_paragraph(paragraph: Paragraph, run, *, role: str) -> None:
    paragraph_format = paragraph.paragraph_format
    if role == "name":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(1)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Aptos"
    elif role == "contact":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(2)
        run.font.size = Pt(9)
        run.font.name = "Aptos"
    elif role == "section_heading":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.space_before = Pt(8)
        paragraph_format.space_after = Pt(2)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = "Aptos"
    elif role == "item_heading":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.space_before = Pt(2)
        paragraph_format.space_after = Pt(0)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Aptos"
    elif role == "bullet":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.left_indent = Inches(0.25)
        paragraph_format.first_line_indent = Inches(-0.18)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        run.font.size = Pt(10)
        run.font.name = "Aptos"
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        run.font.size = Pt(10)
        run.font.name = "Aptos"
    paragraph_format.keep_together = True
    paragraph_format.keep_with_next = role in {"section_heading", "name", "contact"}


def _is_bullet_paragraph(paragraph: Paragraph) -> bool:
    ppr = paragraph._p.pPr
    return bool(ppr is not None and ppr.numPr is not None)


def _is_item_heading_line(line: SourceLine, previous: SourceLine | None, next_line: SourceLine | None) -> bool:
    text = line.text.strip()
    if "heading" in line.style_name.lower() and not _looks_like_heading(text):
        return True
    if _is_role_or_date_line(line):
        return False
    if ACTION_START_RE.search(text):
        return False
    if has_email(text) or has_phone(text):
        return False
    if "—" in text or " - " in text:
        return True
    if next_line and _is_role_or_date_line(next_line):
        return True
    return previous is None and len(text.split()) <= 8


def _is_role_or_date_line(line: SourceLine) -> bool:
    text = line.text.strip()
    return bool((DATE_RE.search(text) or "|" in text) and ROLE_SIGNAL_RE.search(text))


def _looks_like_role_fragment(text: str) -> bool:
    cleaned = text.strip()
    lowered = normalize_for_match(cleaned)
    if ACTION_START_RE.search(cleaned):
        return False
    if has_email(cleaned) or has_phone(cleaned):
        return False
    role_signals = ("developer", "engineer", "analyst", "assistant", "manager", "secretary", "representative", "officer", "lead")
    if not any(signal in lowered for signal in role_signals):
        return False
    if any(separator in cleaned for separator in (" on ", " for ", " at ", " with ", " - ", " — ")):
        return True
    return len(cleaned.split()) <= 16


def _looks_like_bullet_content(text: str) -> bool:
    cleaned = text.strip().lstrip("-*• ")
    return bool(ACTION_START_RE.search(cleaned))


def _polish_bullet(text: str, analysis: AnalysisResult | None) -> str:
    cleaned = text.strip().lstrip("-*• ").strip()
    cleaned = _replace_weak_opening(cleaned)
    cleaned = _rewrite_fragment_sentence(cleaned)
    cleaned = _clean_heading_line(cleaned)
    if analysis:
        cleaned = _add_truthful_context(cleaned, analysis)
    return cleaned.rstrip(".")


def _replace_weak_opening(text: str) -> str:
    replacements = {
        "responsible for": "Owned",
        "worked on": "Contributed to",
        "helped with": "Supported",
        "assisted with": "Supported",
        "participated in": "Contributed to",
    }
    lowered = text.lower()
    for weak, strong in replacements.items():
        if lowered.startswith(weak):
            return strong + text[len(weak) :]
    if lowered.startswith("core ") and " developer on " in lowered:
        return re.sub(r"^core .*? developer on ", "Built frontend functionality for ", text, flags=re.IGNORECASE)
    if lowered.startswith("core ") and " engineer on " in lowered:
        return re.sub(r"^core .*? engineer on ", "Built software for ", text, flags=re.IGNORECASE)
    return text


def _rewrite_fragment_sentence(text: str) -> str:
    lowered = normalize_for_match(text)
    if ACTION_START_RE.search(text):
        return text
    if " developer on " in lowered:
        body = re.split(r"\bdeveloper on\b", text, maxsplit=1, flags=re.IGNORECASE)[1].strip()
        if "frontend" in lowered:
            return f"Built frontend functionality for {body}"
        return f"Developed software for {body}"
    if " engineer on " in lowered:
        body = re.split(r"\bengineer on\b", text, maxsplit=1, flags=re.IGNORECASE)[1].strip()
        return f"Developed software for {body}"
    if any(term in lowered for term in ("dashboard", "platform", "api", "interface", "workflow", "tool")):
        return f"Built {text[0].lower() + text[1:]}"
    if any(term in lowered for term in ("data", "analysis", "analytics", "research", "portfolio", "market")):
        return f"Analyzed {text[0].lower() + text[1:]}"
    if any(term in lowered for term in ("collaborat", "stakeholder", "partnership", "communication")):
        return f"Collaborated on {text[0].lower() + text[1:]}"
    if any(term in lowered for term in ("test", "qa", "validate", "quality")):
        return f"Validated {text[0].lower() + text[1:]}"
    if any(term in lowered for term in ("automate", "script", "workflow")):
        return f"Automated {text[0].lower() + text[1:]}"
    return text


def _add_truthful_context(text: str, analysis: AnalysisResult) -> str:
    return text


def _clean_heading_line(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().strip("•-* ")


def _remove_date_suffix(text: str) -> str:
    return re.split(r"\s+\|\s+|\s{2,}", text, maxsplit=1)[0].strip()


def _compact_context(text: str) -> str:
    cleaned = text.strip().rstrip(".")
    if cleaned.lower().startswith("core ") and " developer on " in cleaned.lower():
        cleaned = re.sub(r"^core .*? developer on ", "developing ", cleaned, flags=re.IGNORECASE)
    if len(cleaned) > 140:
        cleaned = cleaned[:137].rsplit(" ", 1)[0] + "..."
    if not cleaned.lower().startswith(("building", "developing", "working", "supporting", "delivering")):
        cleaned = "developing " + cleaned[0].lower() + cleaned[1:]
    return cleaned


def _strength_terms(analysis: AnalysisResult, resume_text: str) -> list[str]:
    terms = []
    for values in _resume_skill_matches(resume_text, analysis).values():
        terms.extend(values)
    if not terms:
        terms.extend(analysis.matched_keywords[:6])
    return [_title_skill(term) for term in _dedupe(terms)]


def _domain_terms(resume_text: str) -> list[str]:
    haystack = resume_text.lower()
    terms = [term for term in FINANCE_TERMS if term in haystack]
    terms.extend(term for term in TECH_TERMS if term in haystack)
    return [_title_skill(term) for term in _dedupe(terms)]


def _build_interest_line(job_text: str, resume_text: str) -> str:
    terms = _domain_terms(resume_text)
    if not terms:
        return ""
    return " • ".join(terms[:8])


def _research_lines_from_resume(resume_text: str) -> list[str]:
    lines = []
    for line in extract_lines(resume_text):
        lowered = line.lower()
        if any(signal in lowered for signal in ("epq", "dissertation", "thesis", "capstone", "research")):
            lines.append(_clean_research_line(line))
    return _dedupe(lines)[:3]


def _clean_research_line(line: str) -> str:
    cleaned = _clean_heading_line(line)
    if ":" in cleaned:
        left, right = cleaned.split(":", 1)
        return f"{right.strip()} — {left.strip()}"
    return cleaned


def _clean_role(role: str | None) -> str | None:
    if not role:
        return None
    cleaned = re.split(r"\busing\b|\bwith\b|[,.;]", role, maxsplit=1, flags=re.IGNORECASE)[0]
    return re.sub(r"\s+", " ", cleaned).strip(" .")


def _role_with_article(role: str) -> str:
    if re.match(r"^(?:a|an|the)\s+", role, flags=re.IGNORECASE):
        return role
    article = "an" if role[:1].lower() in {"a", "e", "i", "o", "u"} else "a"
    return f"{article} {role}"


def _title_skill(value: str) -> str:
    special = {
        "api": "API",
        "aws": "AWS",
        "c#": "C#",
        "c++": "C++",
        "css": "CSS",
        "gcp": "GCP",
        "git": "Git",
        "html": "HTML",
        "javascript": "JavaScript",
        "machine learning": "Machine Learning",
        "numpy": "NumPy",
        "pandas": "Pandas",
        "python": "Python",
        "react": "React",
        "rest api": "REST APIs",
        "sql": "SQL",
        "typescript": "TypeScript",
        "ui": "UI",
        "ui development": "UI Development",
        "ux": "UX",
        "ux-focused design": "UX-Focused Design",
        "vs code": "VS Code",
        "autohotkey": "AutoHotkey",
        "visualization": "Visualisation",
    }
    normalized = value.strip().lower()
    return special.get(normalized, " ".join(part.capitalize() for part in value.strip().split()))


def _join_human(values: list[str]) -> str:
    if not values:
        return ""
    if len(values) == 1:
        return values[0]
    if len(values) == 2:
        return f"{values[0]} and {values[1]}"
    return ", ".join(values[:-1]) + f" and {values[-1]}"


def _dedupe(values: list[str]) -> list[str]:
    seen = set()
    result = []
    for value in values:
        normalized = value.strip().lower()
        if value and normalized not in seen:
            result.append(value)
            seen.add(normalized)
    return result


def _dedupe_sentences(sentences: list[str]) -> list[str]:
    cleaned = []
    for sentence in sentences:
        sentence = re.sub(r"\s+", " ", sentence).strip()
        if sentence and not sentence.endswith("."):
            sentence += "."
        cleaned.append(sentence)
    return _dedupe(cleaned)
