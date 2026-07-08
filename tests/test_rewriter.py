from __future__ import annotations

from io import BytesIO
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ats_resume_coach.analyzer import ResumeAnalyzer
from ats_resume_coach.rewriter import build_resume_draft


class ResumeRewriterTest(unittest.TestCase):
    def test_builds_docx_draft(self) -> None:
        job = "Software engineering intern using Python FastAPI SQL Git testing communication."
        resume = """
        Jane Jane@example.com github.com/jane
        Education
        BS Computer Science
        Projects
        - Built an API with Python and SQL that processed 1000 records.
        """

        analysis = ResumeAnalyzer().analyze(job, resume)
        draft = build_resume_draft(job, resume, analysis)

        self.assertEqual(draft.content_type, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertTrue(draft.data.startswith(b"PK"))
        self.assertIn("PROFILE", draft.preview_text)

        rewritten = Document(BytesIO(draft.data))
        self.assertEqual(rewritten.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertTrue(rewritten.paragraphs[2].runs[0].bold)
        self.assertTrue(rewritten.paragraphs[2].runs[0].font.size.pt >= 10)
        bullet_paragraph = next(paragraph for paragraph in rewritten.paragraphs if "processed 1000 records" in paragraph.text)
        self.assertIsNotNone(bullet_paragraph.paragraph_format.left_indent)
        self.assertIsNotNone(bullet_paragraph.paragraph_format.first_line_indent)

    def test_docx_rewrite_preserves_source_document(self) -> None:
        job = "Software engineering intern using Python FastAPI SQL Git testing communication."
        source_doc = Document()
        title = source_doc.add_paragraph("Jane Doe")
        title.style = "Title"
        source_doc.add_paragraph("jane@example.com | github.com/jane")
        source_doc.add_paragraph("Profile", style="Heading 1")
        source_doc.add_paragraph("Hard-working student seeking an opportunity.")
        source_doc.add_paragraph("Experience", style="Heading 1")
        source_doc.add_paragraph(
            "Built the original portfolio API with Python and SQL for 1000 records.",
            style="List Bullet",
        )
        source_doc.add_paragraph("Education", style="Heading 1")
        source_doc.add_paragraph("BS Computer Science, University of Leeds")
        source_doc.add_paragraph("Technical Skills", style="Heading 1")
        source_doc.add_paragraph("Programming: Python, SQL Tools: Git")
        source_doc.add_paragraph("Professional Interests", style="Heading 1")
        source_doc.add_paragraph("software engineering")
        source_buffer = BytesIO()
        source_doc.save(source_buffer)

        resume = "\n".join(paragraph.text for paragraph in source_doc.paragraphs)
        analysis = ResumeAnalyzer().analyze(job, resume)
        draft = build_resume_draft(
            job,
            resume,
            analysis,
            source_docx=source_buffer.getvalue(),
            source_filename="jane_resume.docx",
        )
        rewritten = Document(BytesIO(draft.data))
        rewritten_text = "\n".join(paragraph.text for paragraph in rewritten.paragraphs)

        self.assertEqual(draft.filename, "jane_resume_tailored.docx")
        self.assertIn("Jane Doe", rewritten_text)
        self.assertIn("Built the original portfolio API", rewritten_text)
        self.assertIn("BS Computer Science, University of Leeds", rewritten_text)
        self.assertIn("PROFILE", rewritten_text)
        self.assertNotIn("Hard-working student seeking an opportunity.", rewritten_text)
        self.assertNotIn("Tailored Resume Draft", rewritten_text)
        self.assertEqual(rewritten.paragraphs[0].style.name, "Title")
        self.assertIn("BS Computer Science student at University of Leeds", rewritten_text)
        headings = [paragraph.text for paragraph in rewritten.paragraphs if paragraph.text.isupper()]
        self.assertLess(headings.index("EDUCATION"), headings.index("EXPERIENCE"))
        self.assertLess(headings.index("TECHNICAL SKILLS"), headings.index("PROFESSIONAL INTERESTS"))

    def test_text_rewrite_keeps_all_source_sections(self) -> None:
        job = "Data analyst intern using Python SQL Excel communication visualization."
        resume = """
        Alex alex@example.com
        Experience
        - Built dashboard one with Python and SQL.
        - Built dashboard two with Excel and visualization.
        Projects
        - Analyzed customer data with SQL and Python.
        Education
        BS Data Science
        Awards
        Dean's List
        """

        analysis = ResumeAnalyzer().analyze(job, resume)
        draft = build_resume_draft(job, resume, analysis)
        rewritten = Document(BytesIO(draft.data))
        rewritten_text = "\n".join(paragraph.text for paragraph in rewritten.paragraphs)

        self.assertIn("Built dashboard two with Excel", rewritten_text)
        self.assertIn("Analyzed customer data with SQL", rewritten_text)
        self.assertIn("Dean's List", rewritten_text)
        self.assertNotIn("Tailored Resume Draft", rewritten_text)

    def test_docx_rewrite_strengthens_summary_without_mangling_role_lines(self) -> None:
        job = "Commodities Portfolio Manager using Python portfolio construction risk analysis communication."
        source_doc = Document()
        source_doc.add_paragraph("Mariia Luksha", style="Title")
        source_doc.add_paragraph("Leeds, United Kingdom")
        source_doc.add_paragraph("Profile", style="Heading 1")
        source_doc.add_paragraph("Computer Science student interested in finance and technology.")
        source_doc.add_paragraph("Experience", style="Heading 1")
        source_doc.add_paragraph(
            "Core frontend developer on a real-time trading platform used by traders across Binance, Bybit and Hyperliquid."
        )
        source_doc.add_paragraph("Education", style="Heading 1")
        source_doc.add_paragraph("MEng Computer Science (Artificial Intelligence), University of Leeds")
        source_doc.add_paragraph("Technical Skills", style="Heading 1")
        source_doc.add_paragraph("Programming: Python, TypeScript, SQL")
        source_buffer = BytesIO()
        source_doc.save(source_buffer)

        resume = "\n".join(paragraph.text for paragraph in source_doc.paragraphs)
        analysis = ResumeAnalyzer().analyze(job, resume)
        draft = build_resume_draft(
            job,
            resume,
            analysis,
            source_docx=source_buffer.getvalue(),
            source_filename="mariia_resume.docx",
        )
        rewritten = Document(BytesIO(draft.data))
        rewritten_text = "\n".join(paragraph.text for paragraph in rewritten.paragraphs)

        self.assertIn("commercial software engineering experience", rewritten_text)
        self.assertIn("trading technology for financial markets", rewritten_text)
        self.assertIn("Python", rewritten_text)
        self.assertIn("Core frontend developer on a real-time trading platform used by traders across Binance, Bybit and Hyperliquid.", rewritten_text)
        self.assertNotIn("Delivered core frontend developer", rewritten_text)

    def test_analysis_contains_edit_plan(self) -> None:
        job = "Data analyst intern using Python SQL Excel communication visualization."
        resume = """
        Alex alex@example.com
        Education
        BS Data Science
        """

        analysis = ResumeAnalyzer().analyze(job, resume)
        actions = [item.action for item in analysis.edit_plan]
        sections = [item.section for item in analysis.edit_plan]

        self.assertIn("insert", actions)
        self.assertIn("Projects", sections)
        self.assertIn("Skills", sections)


if __name__ == "__main__":
    unittest.main()
