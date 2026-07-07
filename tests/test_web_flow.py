from __future__ import annotations

from io import BytesIO
import re
import unittest

from docx import Document
from fastapi.testclient import TestClient

from ats_resume_coach.api import create_app


JOB_TEXT = "Software engineering intern using Python FastAPI SQL Git testing communication."
RESUME_TEXT = """
Jane jane@example.com github.com/jane
Education
BS Computer Science
Projects
- Built an API with Python and SQL that processed 1000 records.
"""


class WebFlowTest(unittest.TestCase):
    def setUp(self) -> None:
        self.client = TestClient(create_app())

    def test_analyze_preserves_pasted_text_for_rewrite(self) -> None:
        response = self.client.post(
            "/analyze",
            data={"job_text": JOB_TEXT, "resume_text": RESUME_TEXT},
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn(JOB_TEXT, response.text)
        self.assertIn("Jane jane@example.com", response.text)

    def test_analyze_preserves_uploaded_text_for_rewrite(self) -> None:
        response = self.client.post(
            "/analyze",
            files={
                "job_file": ("job.txt", JOB_TEXT.encode(), "text/plain"),
                "resume_file": ("resume.txt", RESUME_TEXT.encode(), "text/plain"),
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn(JOB_TEXT, response.text)
        self.assertIn("Jane jane@example.com", response.text)

    def test_rewrite_after_analyze_reuses_uploaded_docx_source(self) -> None:
        source_doc = Document()
        title = source_doc.add_paragraph("Jane Doe")
        title.style = "Title"
        source_doc.add_paragraph("jane@example.com | github.com/jane")
        source_doc.add_paragraph("Experience", style="Heading 1")
        source_doc.add_paragraph("Built the original trading tool with Python and SQL.", style="List Bullet")
        source_buffer = BytesIO()
        source_doc.save(source_buffer)

        analyze_response = self.client.post(
            "/analyze",
            data={"job_text": JOB_TEXT},
            files={
                "resume_file": (
                    "jane_resume.docx",
                    source_buffer.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            },
        )
        token_match = re.search(r'name="resume_source_token" value="([^"]+)"', analyze_response.text)
        self.assertIsNotNone(token_match)

        rewrite_response = self.client.post(
            "/rewrite",
            data={
                "job_text": JOB_TEXT,
                "resume_text": "Jane Doe\nExperience\nBuilt the original trading tool with Python and SQL.",
                "resume_source_token": token_match.group(1),
            },
        )
        rewritten = Document(BytesIO(rewrite_response.content))
        rewritten_text = "\n".join(paragraph.text for paragraph in rewritten.paragraphs)

        self.assertEqual(rewrite_response.status_code, 200)
        self.assertIn("Jane Doe", rewritten_text)
        self.assertIn("Built the original trading tool", rewritten_text)
        self.assertIn("Professional Summary", rewritten_text)
        self.assertNotIn("Tailored Resume Draft", rewritten_text)
        self.assertEqual(rewritten.paragraphs[0].style.name, "Title")

    def test_rewrite_downloads_docx_from_form_text(self) -> None:
        response = self.client.post(
            "/rewrite",
            data={"job_text": JOB_TEXT, "resume_text": RESUME_TEXT},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(response.content.startswith(b"PK"))

    def test_rewrite_validation_error_renders_html_form(self) -> None:
        response = self.client.post("/rewrite", data={})

        self.assertEqual(response.status_code, 400)
        self.assertIn("text/html", response.headers["content-type"])
        self.assertIn("Provide a job URL, job file, or pasted job text.", response.text)


if __name__ == "__main__":
    unittest.main()
